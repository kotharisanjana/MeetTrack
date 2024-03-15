import openai
from dotenv import load_dotenv, find_dotenv
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document
from langchain.llms.openai import OpenAI, OpenAIChat
from langchain.prompts import PromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from aws_utilities import  get_s3_image_bytes, download_file_from_s3, upload_file_to_s3
from image_context import get_images_from_context
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image
import boto3
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os

_ = load_dotenv(find_dotenv())
openai.api_key = os.environ['OPENAI_API_KEY']

s3_doc_path = "s3://bucket_name/path_to_doc"
bucket_name = ""
object_key=""

def generate_textual_components():
    file = download_file() #transcript file
    target_len = 500
    chunk_size = 3000
    chunk_overlap = 200
    with open("file", "r") as f:
        raw_text = f.read()
    # Split the source text
    text_splitter = CharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    texts = text_splitter.split_text(
        raw_text,
    )

    # Create Document objects for the texts
    docs = [Document(page_content=t) for t in texts[:]]

    openaichat = OpenAIChat(temperature=0, model="gpt-3.5-turbo")
    prompt_template = """Act as a professional technical meeting minutes writer.
    Tone: formal
    Format: Technical meeting summary
    Length:  200 ~ 300
    Tasks:
    - highlight action items and owners
    - highlight the agreements
    - Use bullet points if needed
    {text}
    CONCISE SUMMARY IN ENGLISH:"""
    PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])
    refine_template = (
        "Your job is to produce a final summary\n"
        "We have provided an existing summary up to a certain point: {existing_answer}\n"
        "We have the opportunity to refine the existing summary"
        "(only if needed) with some more context below.\n"
        "------------\n"
        "{text}\n"
        "------------\n"
        f"Given the new context, refine the original summary in English within {target_len} words: following the format"
        "Participants: <participants>"
        "Discussed: <Discussed-items>"
        "Follow-up actions: <a-list-of-follow-up-actions-with-owner-names>"
        "If the context isn't useful, return the original summary. Highlight agreements and follow-up actions and owners."
    )
    refine_prompt = PromptTemplate(
        input_variables=["existing_answer", "text"],
        template=refine_template,
    )
    chain = load_summarize_chain(
        openaichat,
        chain_type="refine",
        return_intermediate_steps=True,
        question_prompt=PROMPT,
        refine_prompt=refine_prompt,
    )
    resp = chain({"input_documents": docs}, return_only_outputs=True)
    textual_document = resp["output_text"]

    #Need to extract only summary and pass it instead of entire textual doc
    relevant_image_links = get_images_from_context(textual_document)
    return textual_document, relevant_image_links

def create_doc_with_text_and_images(textual_document, image_keys, bucket_name, s3_doc_path):

    doc = Document()
    doc.add_paragraph(textual_document)
    
    for key in image_keys:
        image_bytes = get_s3_image_bytes(bucket_name, key)
        if image_bytes:
            # Convert bytes to a PIL Image to determine the image size
            image = Image.open(BytesIO(image_bytes))
            width, height = image.size
            # Calculate the image's width and height in inches for the document
            aspect_ratio = height / width
            display_width = min(4.0, width / 72) 
            display_height = display_width * aspect_ratio
            
            # Add image to the document
            image_stream = BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(display_width), height=Inches(display_height))
    
    doc.save(local_file_path)
    upload_file_to_s3(bucket_name,s3_doc_path,local_file_path)

def send_email():
    textual_document,image_keys = generate_textual_components()
    create_doc_with_text_and_images(textual_document, image_keys, bucket_name, s3_doc_path)
    download_file_from_s3(bucket_name, object_key, local_file_path)
    #email config
    recipient_email = "ananya.joshi@sjsu.edu"
    sender_email = "ananya.joshi@sjsu.edu"
    sender_password = "****"
    subject = "Meeting Track Document"
    body = "Please find the attached document."

    # Create MIME multipart message
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    # Attach the body with the msg instance
    msg.attach(MIMEText(body, 'plain'))

    # Open the file to be sent
    with open(local_file_path, "rb") as attachment:
        part = MIMEApplication(attachment.read(), Name=os.path.basename(local_file_path))
    # After the file is closed
    part['Content-Disposition'] = f'attachment; filename="{os.path.basename(local_file_path)}"'
    msg.attach(part)

    # Setup the SMTP server
    try:
        server = smtplib.SMTP('smtp.example.com', 587)  # Use your SMTP server and port
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, msg.as_string())
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print(f"Failed to send email: {e}")



    

