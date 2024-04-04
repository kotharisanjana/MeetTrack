from common.aws_utilities import get_s3_image_bytes, upload_file_to_s3
from database.relational_db import fetch_output_path
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

def create_final_doc(session_data, textual_component, image_url_desc_pairs):
    meeting_id = session_data["meeting_id"]
    local_output_path = session_data["local_output_path"]
    meeting_name = session_data["meeting_name"]
    meeting_date = session_data["meeting_date"]

    header = "Meeting Name: " + meeting_name + "\n" + "Meeting Date: " + meeting_date

    doc = Document()
    doc.add_paragraph(header)
    doc.add_paragraph(textual_component)
    
    for url, desc in image_url_desc_pairs.items():
        image_bytes = get_s3_image_bytes(url)

        if image_bytes:
            # Convert bytes to a PIL Image
            image = Image.open(BytesIO(image_bytes))

            # Calculate the image's width and height in inches for the document
            width, height = image.size
            aspect_ratio = height / width
            display_width = min(4.0, width / 72) 
            display_height = display_width * aspect_ratio

            # Add image to the document
            image_stream = BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(display_width), height=Inches(display_height))
            doc.add_paragraph("Image description: " + desc)
            
    doc.save(local_output_path)
    output_path = fetch_output_path(meeting_id)
    upload_file_to_s3(local_output_path, output_path)